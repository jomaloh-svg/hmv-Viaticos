import sys, json, base64, os, shutil, subprocess
from docx import Document
from docx.oxml.ns import qn
import lxml.etree as etree

WPS_NS = 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape'
W_NS   = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
WP_NS  = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'

# Mapa: nombre del recuadro → significado
# Rectangle 3 = PEN, Rectangle 4 = USD
# Rectangle 5 = Ocasionales, Rectangle 6 = Permanentes, Rectangle 2 = Combustible

def get_shapes(cell):
    """Devuelve dict {nombre: wsp_element} para todos los recuadros flotantes de la celda."""
    shapes = {}
    for para in cell.paragraphs:
        for run in para.runs:
            for anchor in run._r.iter(f'{{{WP_NS}}}anchor'):
                docPr = anchor.find(f'{{{WP_NS}}}docPr')
                name  = docPr.get('name') if docPr is not None else ''
                wsp   = anchor.find(f'.//{{{WPS_NS}}}wsp')
                if wsp is not None:
                    shapes[name] = wsp
    return shapes

def add_x_to_shape(wsp_el):
    """Inserta una X centrada vertical y horizontalmente dentro del recuadro flotante."""
    # 1. Ajustar bodyPr para centrado vertical y sin márgenes internos
    bodyPr = wsp_el.find(f'{{{WPS_NS}}}bodyPr')
    if bodyPr is not None:
        bodyPr.set('anchor',    'ctr')   # centrado vertical
        bodyPr.set('anchorCtr', '1')
        bodyPr.set('lIns',      '0')     # sin margen interno
        bodyPr.set('rIns',      '0')
        bodyPr.set('tIns',      '0')
        bodyPr.set('bIns',      '0')

    # 2. Construir txbx con X centrada
    txbx    = etree.Element(f'{{{WPS_NS}}}txbx')
    content = etree.SubElement(txbx, f'{{{W_NS}}}txbxContent')
    p       = etree.SubElement(content, f'{{{W_NS}}}p')
    pPr     = etree.SubElement(p, f'{{{W_NS}}}pPr')
    jc      = etree.SubElement(pPr, f'{{{W_NS}}}jc')
    jc.set(f'{{{W_NS}}}val', 'center')
    spc     = etree.SubElement(pPr, f'{{{W_NS}}}spacing')
    spc.set(f'{{{W_NS}}}before', '0')
    spc.set(f'{{{W_NS}}}after',  '0')
    r       = etree.SubElement(p, f'{{{W_NS}}}r')
    rPr     = etree.SubElement(r, f'{{{W_NS}}}rPr')
    b       = etree.SubElement(rPr, f'{{{W_NS}}}b')
    sz      = etree.SubElement(rPr, f'{{{W_NS}}}sz')
    sz.set(f'{{{W_NS}}}val', '14')
    t       = etree.SubElement(r, f'{{{W_NS}}}t')
    t.text  = 'X'
    wsp_el.insert(list(wsp_el).index(bodyPr), txbx)

def remove_highlight(run):
    rPr = run._r.find(qn('w:rPr'))
    if rPr is not None:
        for tag in [qn('w:highlight'), qn('w:shd')]:
            el = rPr.find(tag)
            if el is not None:
                rPr.remove(el)

def set_run(run, text):
    run.text = text
    remove_highlight(run)

def numero_en_letras(monto):
    entero   = int(monto)
    centavos = round((monto - entero) * 100)
    unidades = ['','UNO','DOS','TRES','CUATRO','CINCO','SEIS','SIETE','OCHO','NUEVE','DIEZ',
        'ONCE','DOCE','TRECE','CATORCE','QUINCE','DIECISÉIS','DIECISIETE','DIECIOCHO','DIECINUEVE',
        'VEINTE','VEINTIUNO','VEINTIDÓS','VEINTITRÉS','VEINTICUATRO','VEINTICINCO','VEINTISÉIS',
        'VEINTISIETE','VEINTIOCHO','VEINTINUEVE']
    decenas  = ['','','VEINTE','TREINTA','CUARENTA','CINCUENTA','SESENTA','SETENTA','OCHENTA','NOVENTA']
    centenas = ['','CIENTO','DOSCIENTOS','TRESCIENTOS','CUATROCIENTOS','QUINIENTOS',
        'SEISCIENTOS','SETECIENTOS','OCHOCIENTOS','NOVECIENTOS']

    def grupo(n):
        if n == 0: return ''
        if n == 100: return 'CIEN'
        r = ''
        if n >= 100:
            r += centenas[n // 100] + ' '; n %= 100
        if n < 30: r += unidades[n]
        else:
            r += decenas[n // 10]
            if n % 10: r += ' Y ' + unidades[n % 10]
        return r.strip()

    def convertir(n):
        if n == 0: return 'CERO'
        r = ''
        if n >= 1000:
            miles = n // 1000
            r += ('MIL' if miles == 1 else grupo(miles) + ' MIL') + ' '
            n %= 1000
        r += grupo(n)
        return r.strip()

    return f"{convertir(entero)} {str(centavos).zfill(2)}/100 NUEVOS SOLES"

def fmt_date(iso):
    if not iso: return ''
    parts = iso.split('-')
    return f"{parts[2]}/{parts[1]}/{parts[0]}" if len(parts) == 3 else iso

def fill_and_convert(data):
    src = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'plantilla.docx')
    dst = '/tmp/filled.docx'
    shutil.copy(src, dst)
    doc = Document(dst)

    nombre       = data.get('nombre', '')
    dni          = data.get('dni', '')
    proyecto     = data.get('proyecto', '')
    consec       = data.get('consec', '001')
    fecha_elab   = fmt_date(data.get('fechaElab', ''))
    ciudad       = data.get('ciudad', '')
    motivo       = data.get('motivo', '')
    comentarios  = data.get('comentarios', '')
    deductivo    = float(data.get('deductivo', 0))
    deductivo_desc = data.get('deductivoDesc', '')
    fecha_sal    = fmt_date(data.get('fechaSalida', ''))
    fecha_reg    = fmt_date(data.get('fechaRegreso', ''))
    monto        = float(data.get('monto', 0))
    tipo         = data.get('tipo', 'Ocasionales')
    moneda       = data.get('moneda', 'PEN')
    cuenta       = data.get('cuenta', '')
    cci          = data.get('cci', '')
    banco        = data.get('banco', '')
    ticket_de    = data.get('ticketDe', '')
    ticket_a     = data.get('ticketA', '')
    ticket_fecha = fmt_date(data.get('ticketFecha', ''))
    ticket_hora  = data.get('ticketHora', '')
    ticket_de2   = data.get('ticketDe2', '')
    ticket_a2    = data.get('ticketA2', '')
    ticket_fecha2= fmt_date(data.get('ticketFecha2', ''))
    ticket_hora2 = data.get('ticketHora2', '')
    dias         = int(data.get('dias', 0))

    # Texto de comentarios con deductivo integrado
    comentarios_full = comentarios
    if deductivo > 0:
        ded_texto = f"Deductivo: S/. {deductivo:.2f}"
        if deductivo_desc:
            ded_texto += f" — {deductivo_desc}"
        comentarios_full = (comentarios_full + ' | ' + ded_texto).strip(' |')

    monto_letras = numero_en_letras(monto)
    monto_str    = f"{monto:.2f}".split('.')
    monto_ent    = f"{int(monto_str[0]):,}"
    monto_dec    = monto_str[1]

    tables = doc.tables

    # ── TABLE 1: Consecutivo y Fecha elaboración ──────────────────
    t1   = tables[1]
    runs = t1.rows[1].cells[0].paragraphs[0].runs
    if len(runs) > 2: set_run(runs[2], consec)
    if len(runs) > 3: set_run(runs[3], '')

    runs_fe = t1.rows[1].cells[1].paragraphs[0].runs
    dp = fecha_elab.split('/') if fecha_elab else ['','','']
    if len(runs_fe) > 2: set_run(runs_fe[2], dp[0] if len(dp) > 0 else '')
    if len(runs_fe) > 3: set_run(runs_fe[3], '/')
    if len(runs_fe) > 4: set_run(runs_fe[4], dp[1] if len(dp) > 1 else '')
    if len(runs_fe) > 5: set_run(runs_fe[5], '/' + (dp[2] if len(dp) > 2 else ''))

    # ── TABLE 2: Datos del empleado ───────────────────────────────
    t2 = tables[2]

    runs = t2.rows[0].cells[0].paragraphs[0].runs
    if len(runs) > 2: set_run(runs[2], nombre)

    runs = t2.rows[1].cells[0].paragraphs[0].runs
    if len(runs) > 2: set_run(runs[2], dni)

    runs = t2.rows[2].cells[0].paragraphs[0].runs
    if len(runs) > 3: set_run(runs[3], proyecto)

    runs = t2.rows[3].cells[0].paragraphs[0].runs
    if len(runs) > 2: set_run(runs[2], ciudad)

    runs = t2.rows[4].cells[0].paragraphs[0].runs
    if len(runs) > 2: set_run(runs[2], motivo)

    runs = t2.rows[5].cells[0].paragraphs[0].runs
    if len(runs) > 2: set_run(runs[2], comentarios_full)

    # Fechas
    runs = t2.rows[6].cells[0].paragraphs[0].runs
    sal  = fecha_sal.split('/') if fecha_sal else ['','','']
    reg  = fecha_reg.split('/') if fecha_reg else ['','','']
    if len(runs) > 7:  set_run(runs[7],  sal[0] if len(sal) > 0 else '')
    if len(runs) > 8:  set_run(runs[8],  '/')
    if len(runs) > 9:  set_run(runs[9],  sal[1] if len(sal) > 1 else '')
    if len(runs) > 10: set_run(runs[10], '/' + (sal[2] if len(sal) > 2 else ''))
    if len(runs) > 17: set_run(runs[17], reg[0] if len(reg) > 0 else '')
    if len(runs) > 18: set_run(runs[18], '/')
    if len(runs) > 19: set_run(runs[19], reg[1] if len(reg) > 1 else '')
    if len(runs) > 20: set_run(runs[20], '/' + (reg[2] if len(reg) > 2 else ''))

    # ── Monto ─────────────────────────────────────────────────────
    cell  = t2.rows[7].cells[0]
    runs0 = cell.paragraphs[0].runs

    if deductivo > 0:
        # Monto bruto (antes del deductivo)
        monto_bruto     = monto + deductivo
        bruto_str       = f"{monto_bruto:.2f}".split('.')
        bruto_ent       = f"{int(bruto_str[0]):,}"
        bruto_dec       = bruto_str[1]
        ded_ent         = f"{int(deductivo):,}"
        ded_dec         = f"{deductivo:.2f}".split('.')[1]

        # Formato: S/. 510.00 − S/. 170.00 = S/. 340.00 (LETRAS)
        ecuacion = (
            f"S/. {bruto_ent}.{bruto_dec} "
            f"\u2212 S/. {ded_ent}.{ded_dec} "
            f"= S/. {monto_ent}.{monto_dec} "
            f"({monto_letras})"
        )
        # Colocar toda la ecuación en run1, vaciar los demás runs de P0
        if len(runs0) > 0: set_run(runs0[0], 'Solicito la suma de:  ')
        if len(runs0) > 1: set_run(runs0[1], ecuacion)
        for i in range(2, len(runs0)):
            set_run(runs0[i], '')
    else:
        # Sin deductivo: formato original S/ X,XXX.00 (LETRAS)
        if len(runs0) > 1: set_run(runs0[1], 'S/ ')
        if len(runs0) > 2: set_run(runs0[2], monto_ent)
        if len(runs0) > 3: set_run(runs0[3], '.')
        if len(runs0) > 4: set_run(runs0[4], monto_dec)
        if len(runs0) > 5: set_run(runs0[5], '')
        if len(runs0) > 6: set_run(runs0[6], '')
        if len(runs0) > 7: set_run(runs0[7], ' (')
        if len(runs0) > 8: set_run(runs0[8], monto_letras)
        if len(runs0) > 9: set_run(runs0[9], ')')

    # ── Eliminar recuadros flotantes ─────────────────────────────
    WP_NS_  = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
    MC_NS   = 'http://schemas.openxmlformats.org/markup-compatibility/2006'
    for pi in [1, 4]:
        para = cell.paragraphs[pi]
        for run in para.runs:
            # Remove AlternateContent (wraps the anchor drawing)
            for ac in run._r.findall(f'{{{MC_NS}}}AlternateContent'):
                run._r.remove(ac)
            # Remove any remaining anchor
            for anc in run._r.findall(f'.//{{{WP_NS_}}}anchor'):
                anc.getparent().remove(anc)

    # ── Marcar X en texto según selección ────────────────────────
    # Font size original = 127000 EMU = 10pt → use 152400 = 12pt (1pt larger)
    X_SIZE = '24'   # half-points: 24 = 12pt

    def make_x_run(para_obj):
        """Return a new bold run with X at larger size, cloned from run3 formatting."""
        import copy
        from docx.oxml.ns import qn as qn2
        base = para_obj.runs[3]._r  # use run3 as format reference
        new_r = copy.deepcopy(base)
        # Clear text
        for t in new_r.findall(qn2('w:t')):
            new_r.remove(t)
        # Set rPr: bold + size
        rPr = new_r.find(qn2('w:rPr'))
        if rPr is None:
            rPr = etree.SubElement(new_r, qn2('w:rPr'))
        b = rPr.find(qn2('w:b'))
        if b is None:
            b = etree.SubElement(rPr, qn2('w:b'))
        sz = rPr.find(qn2('w:sz'))
        if sz is None:
            sz = etree.SubElement(rPr, qn2('w:sz'))
        sz.set(qn2('w:val'), X_SIZE)
        szCs = rPr.find(qn2('w:szCs'))
        if szCs is None:
            szCs = etree.SubElement(rPr, qn2('w:szCs'))
        szCs.set(qn2('w:val'), X_SIZE)
        # Add text X
        t_el = etree.SubElement(new_r, qn2('w:t'))
        t_el.text = 'X'
        return new_r

    # ── Marcar X después de cada opción: "Opción: X" ────────────
    # Format: texto de opción + ': ' + X en negrita tamaño 12pt

    def insert_colon_x_after(run_obj, para_obj):
        """Inserta ': X' después del run indicado. X en negrita y 12pt."""
        import copy
        from docx.oxml.ns import qn as qn2
        XML_SP = '{http://www.w3.org/XML/1998/namespace}space'
        # Run con ': '
        sep_r = copy.deepcopy(run_obj._r)
        for t in sep_r.findall(qn2('w:t')): sep_r.remove(t)
        rPr = sep_r.find(qn2('w:rPr'))
        if rPr is not None:
            for tag in [qn2('w:b'), qn2('w:bCs')]:
                el = rPr.find(tag)
                if el is not None: rPr.remove(el)
        t_sep = etree.SubElement(sep_r, qn2('w:t'))
        t_sep.text = ': '
        t_sep.set(XML_SP, 'preserve')
        run_obj._r.addnext(sep_r)
        # Run con X en negrita 12pt
        x_r = copy.deepcopy(run_obj._r)
        for t in x_r.findall(qn2('w:t')): x_r.remove(t)
        rPr2 = x_r.find(qn2('w:rPr'))
        if rPr2 is None:
            rPr2 = etree.SubElement(x_r, qn2('w:rPr'))
        # Bold
        if rPr2.find(qn2('w:b')) is None:
            etree.SubElement(rPr2, qn2('w:b'))
        if rPr2.find(qn2('w:bCs')) is None:
            etree.SubElement(rPr2, qn2('w:bCs'))
        # Size 12pt = 24 half-points
        sz = rPr2.find(qn2('w:sz'))
        if sz is None: sz = etree.SubElement(rPr2, qn2('w:sz'))
        sz.set(qn2('w:val'), '24')
        szCs = rPr2.find(qn2('w:szCs'))
        if szCs is None: szCs = etree.SubElement(rPr2, qn2('w:szCs'))
        szCs.set(qn2('w:val'), '24')
        t_x = etree.SubElement(x_r, qn2('w:t'))
        t_x.text = 'X'
        sep_r.addnext(x_r)

    # P1 — Moneda: after 'PEN)' (run8) or '(USD)' (run15)
    runs1 = cell.paragraphs[1].runs
    runs1[10].text = '    '   # clear original X residual
    if moneda == 'PEN':
        insert_colon_x_after(runs1[8], cell.paragraphs[1])   # after 'PEN)'
    else:
        insert_colon_x_after(runs1[15], cell.paragraphs[1])  # after '(USD)'

    # P4 — Tipo viático: reescritura completa sin tabs, una sola línea
    from docx.oxml import OxmlElement as OXE
    import copy as _copy
    p4   = cell.paragraphs[4]._p
    pPr4 = p4.find(qn('w:pPr'))

    def mk_run(text, bold=False, sz='20'):
        r  = OXE('w:r'); rP = OXE('w:rPr')
        s  = OXE('w:sz');   s.set(qn('w:val'), sz);  rP.append(s)
        sC = OXE('w:szCs'); sC.set(qn('w:val'), sz); rP.append(sC)
        if bold: rP.append(OXE('w:b')); rP.append(OXE('w:bCs'))
        r.append(rP)
        t = OXE('w:t'); t.text = text
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r.append(t); return r

    # Remove all w:r from P4
    for ch in list(p4):
        if ch.tag == qn('w:r'): p4.remove(ch)

    # Remove tabs and firstLine indent
    for tag in [qn('w:tabs'), qn('w:ind')]:
        el = pPr4.find(tag)
        if el is not None: pPr4.remove(el)
    ni = OXE('w:ind'); ni.set(qn('w:left'), '554'); pPr4.append(ni)

    # One line: label + ': X' for selected tipo, plain text for others
    p4.append(mk_run('Viáticos Ocasionales'))
    if tipo == 'Ocasionales': p4.append(mk_run(': ', sz='20')); p4.append(mk_run('X', bold=True, sz='24'))
    p4.append(mk_run('      '))
    p4.append(mk_run('Viáticos Permanentes'))
    if tipo == 'Permanentes': p4.append(mk_run(': ', sz='20')); p4.append(mk_run('X', bold=True, sz='24'))
    p4.append(mk_run('      '))
    p4.append(mk_run('Combustible'))
    if tipo == 'Combustible': p4.append(mk_run(': ', sz='20')); p4.append(mk_run('X', bold=True, sz='24'))

    # Separate paragraph for "Y que serán..."
    new_p4b = OXE('w:p')
    new_pPr4 = _copy.deepcopy(pPr4)
    ni2 = new_pPr4.find(qn('w:ind'))
    if ni2 is not None: new_pPr4.remove(ni2)
    new_p4b.append(new_pPr4)
    new_p4b.append(mk_run('Y que serán sustentados con las facturas o boletas de pago a mi regreso.'))
    p4.addnext(new_p4b)

    # ── Datos bancarios ───────────────────────────────────────────
    for run in t2.rows[9].cells[0].paragraphs[1].runs:
        if 'X' in run.text: set_run(run, cuenta)
    for run in t2.rows[9].cells[1].paragraphs[1].runs:
        if 'X' in run.text: set_run(run, cci)
    for run in t2.rows[9].cells[2].paragraphs[1].runs:
        if 'X' in run.text: set_run(run, banco)

    # ── TABLE 3: Tiquete aéreo ────────────────────────────────────
    if ticket_de or ticket_a or ticket_de2 or ticket_a2:
        t3 = tables[3]
        def set_cell_text(c, text):
            p = c.paragraphs[0]
            if not p.runs: p.add_run(text)
            else: p.runs[0].text = text
        # Vuelo 1 — fila 2
        if len(t3.rows) > 2:
            row = t3.rows[2]
            if len(row.cells) > 1: set_cell_text(row.cells[1], ticket_de)
            if len(row.cells) > 2: set_cell_text(row.cells[2], ticket_a)
            if len(row.cells) > 3: set_cell_text(row.cells[3], ticket_fecha)
            if len(row.cells) > 4: set_cell_text(row.cells[4], ticket_hora)
        # Vuelo 2 — fila 3
        if (ticket_de2 or ticket_a2) and len(t3.rows) > 3:
            row2 = t3.rows[3]
            if len(row2.cells) > 1: set_cell_text(row2.cells[1], ticket_de2)
            if len(row2.cells) > 2: set_cell_text(row2.cells[2], ticket_a2)
            if len(row2.cells) > 3: set_cell_text(row2.cells[3], ticket_fecha2)
            if len(row2.cells) > 4: set_cell_text(row2.cells[4], ticket_hora2)

    doc.save(dst)

    result = subprocess.run(
        ['soffice', '--headless', '--convert-to', 'pdf', '--outdir', '/tmp/', dst],
        capture_output=True, text=True, timeout=60
    )
    pdf_path = '/tmp/filled.pdf'
    if not os.path.exists(pdf_path):
        raise Exception(f"PDF conversion failed: {result.stderr or result.stdout}")

    with open(pdf_path, 'rb') as f:
        b64 = base64.b64encode(f.read()).decode('utf-8')

    os.remove(pdf_path)
    os.remove(dst)
    return b64

if __name__ == '__main__':
    data = json.loads(sys.argv[1])
    print(fill_and_convert(data))
